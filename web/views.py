import os
import json
from io import BytesIO
import traceback
from collections import defaultdict
from PIL import Image,ImageEnhance, ImageOps
import pandas as pd
# import cv2
# import numpy as np
from langdetect import detect, DetectorFactory
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from PyPDF2 import PdfReader # While imported, it's not used in extract_text_from_file
import pdfplumber
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
DetectorFactory.seed = 0

def summarizer_page(request):
    """
    Trang tóm tắt văn bản.
    """
    return render(request, 'summarizer.html')

@csrf_exempt
def summarize_text(request):
    """
    Hàm xử lý yêu cầu tóm tắt văn bản.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Phương thức không được hỗ trợ'}, status=405)

    try:
        body = json.loads(request.body)
        content = body.get('content', '').strip()
        detected_language = body.get('detected_language')

        if not content:
            return JsonResponse({'error': 'Không có nội dung để tóm tắt'}, status=400)
        if not detected_language:
            try:
                detected_language = detect(content)
            except:
                detected_language = 'unknown'

        language_map = {
                'vi': 'tiếng Việt',
                'en': 'English',
                'fr': 'French',
                'de': 'German',
                'ja': 'Japanese',
                'zh-cn': 'Chinese',
                'unknown': 'ngôn ngữ gốc'
            }
        language_name = language_map.get(detected_language, 'ngôn ngữ gốc')

        prompt = f"""Bạn là một trợ lý AI thông minh. Hãy tóm tắt văn bản sau bằng đúng ngôn ngữ gốc của nó, hiện đang được phát hiện là **{language_name}**.
            Yêu cầu:
            - Tóm tắt phải sử dụng đúng ngôn ngữ của văn bản gốc: {language_name}.
            - Chỉ nêu các ý chính, loại bỏ chi tiết rườm rà.
            - Tóm tắt tối đa trong 1 đến 2 trang word
            - Tổng hợp tất cả số liệu 
            Văn bản cần tóm tắt:
            {content}
            """

        # Lấy API key từ biến môi trường
        api_key = os.getenv("GEMINI_API_KEY")

        # Kiểm tra xem API key có tồn tại không
        if not api_key:
            return JsonResponse({'error': 'Không tìm thấy API key'}, status=500)

        genai.configure(api_key=api_key)

        model = genai.GenerativeModel('gemini-1.5-flash')

        response = model.generate_content(prompt)

        return JsonResponse({'summary': response.text})

    except Exception as e:
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)

def export_summary(request):
    """
    Hàm xử lý yêu cầu xuất bản tóm tắt.
    """
    if request.method == "POST":
        summary_text = request.POST.get("summary", "")
        export_format = request.POST.get("format", "")

        if export_format == "docx":
            buffer = BytesIO()
            doc = Document()

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Đặt font đúng với ngôn ngữ tiếng Việt (phải đặt cả name và element)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            for line in summary_text.strip().split("\n"):
                if line.strip():
                    para = doc.add_paragraph(line.strip())

                    # Cài đặt spacing cho từng đoạn
                    para_format = para.paragraph_format
                    para_format.space_before = Pt(0)
                    para_format.space_after = Pt(6)
                    para_format.line_spacing = 1.15

            doc.save(buffer)
            buffer.seek(0)
            return HttpResponse(
                buffer,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="tom_tat.docx"'}
            )

    return HttpResponse("Yêu cầu không hợp lệ", status=400)

@csrf_exempt
def extract_text_from_file(request):
    """
    Nhận file PDF hoặc DOCX, trích xuất văn bản và dữ liệu từ bảng.
    """
    if request.method != "POST":
        return JsonResponse({'error': 'Phương thức không được hỗ trợ'}, status=405)

    try:
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return JsonResponse({'error': 'Không có file nào được tải lên'}, status=400)
        if 'summary_result' in request.session:
            del request.session['summary_result']

        filename = uploaded_file.name.lower()
        content = ""

        if filename.endswith('.pdf'):
            with pdfplumber.open(uploaded_file) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_content = ""

                    # 1. Thử dùng pdfplumber để trích xuất văn bản
                    words = page.extract_words()
                    if words:
                        lines = defaultdict(list)
                        for word in words:
                            top_rounded = round(word['top'] / 2) * 2
                            lines[top_rounded].append((word['x0'], word['text']))

                        sorted_lines = sorted(lines.items(), key=lambda x: x[0])
                        all_lines = []
                        for _, line_words in sorted_lines:
                            sorted_words = sorted(line_words, key=lambda x: x[0])
                            line_text = " ".join(w for _, w in sorted_words)
                            all_lines.append(line_text.strip())

                        paragraph = " ".join(all_lines)
                        page_content += f"[{i+1}]\n{paragraph}\n\n"
                        print(f"DEBUG: Page {i+1} - pdfplumber extracted text.")
                    else:
                        # 2. Nếu không có text, dùng OCR
                        print(f"DEBUG: Page {i+1} - pdfplumber found no text, attempting OCR.")
                        try:
                            pil_image = page.to_image(resolution=300).original.convert("RGB")
                            ocr_text = extract_text_with_ocr_data(pil_image)
                            if ocr_text and ocr_text.strip():
                                page_content += f"[{i+1} - OCR]\n{ocr_text.strip()}\n\n"
                                print(f"DEBUG: Page {i+1} - OCR success, extracted {len(ocr_text.strip().split())} words.")
                            else:
                                print(f"DEBUG: Page {i+1} - OCR returned empty text.")
                        except Exception as ocr_e:
                            print(f"ERROR: Page {i+1} - OCR failed: {ocr_e}")
                            page_content += f"[{i+1} - OCR Error] Could not extract text.\n\n"

                    # 3. Trích xuất bảng nếu có
                    tables = page.extract_tables()
                    if tables:
                        print(f"DEBUG: Page {i+1} - Tables found.")
                        for table in tables:
                            for row in table:
                                row_text = "\t".join(cell if cell else "" for cell in row)
                                page_content += row_text + "\n"
                            page_content += "\n"
                    else:
                        print(f"DEBUG: Page {i+1} - No tables found.")

                    content += page_content

        elif filename.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                content += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    content += "\t".join(row_data) + "\n"

        else:
            return JsonResponse({'error': 'Chỉ hỗ trợ các file PDF hoặc DOCX'}, status=400)

        if not content.strip():
            return JsonResponse({'error': 'Không thể trích xuất nội dung từ tệp. Đảm bảo tệp không trống hoặc bị lỗi hình ảnh.'}, status=400)

        return JsonResponse({'content': content})

    except Exception as e:
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)

def extract_text_with_ocr_data(pil_image, lang="eng+vie"):
    """
    Sử dụng pytesseract để OCR ảnh, giữ cấu trúc bảng bằng cách nhóm dòng và phân tách bằng tab.
    """
    try:
        processed_image = preprocess_image_for_ocr(pil_image)

        custom_config = "--psm 6"  # layout OCR
        ocr_df = pytesseract.image_to_data(
            processed_image,
            lang=lang,
            output_type=pytesseract.Output.DATAFRAME,
            config=custom_config
        )

        ocr_df = ocr_df[ocr_df.text.notnull() & (ocr_df.text.str.strip() != '')]
        if ocr_df.empty:
            return None

        # Gom nhóm theo dòng
        grouped = ocr_df.groupby(['block_num', 'par_num', 'line_num'])
        lines = []
        for _, group in grouped:
            words = group.sort_values(by='left')
            line = "\t".join(words['text'].str.strip())
            lines.append(line)

        return "\n".join(lines)

    except Exception as e:
        print(f"ERROR in OCR processing: {e}")
        return None

def preprocess_image_for_ocr(pil_image):
    """
    Tiền xử lý ảnh: grayscale, threshold, tăng contrast và sharpness.
    """
    gray = pil_image.convert("L")
    binary = gray.point(lambda x: 0 if x < 180 else 255, '1')  # Threshold
    rgb = binary.convert("RGB")
    contrast = ImageEnhance.Contrast(rgb).enhance(1.8)
    sharpened = ImageEnhance.Sharpness(contrast).enhance(2.5)
    return sharpened


def extract_table_like_text(ocr_df):
    """
    Gom nhóm các dòng có tọa độ gần nhau để tạo bảng tạm.
    """
    lines = []
    for _, group in ocr_df.groupby(['block_num', 'par_num', 'line_num']):
        words = group.sort_values(by='left')
        line = []
        for _, row in words.iterrows():
            line.append(row['text'].strip())
        lines.append("\t".join(line))
    return "\n".join(lines)
